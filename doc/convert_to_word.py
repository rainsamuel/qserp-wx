#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
将软著操作手册Markdown转换为Word文档
"""

import os
import re
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_shading(cell, color):
    """设置单元格背景色"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading_elm)

def create_document():
    doc = Document()

    # 设置默认字体
    style = doc.styles['Normal']
    font = style.font
    font.name = '宋体'
    font.size = Pt(12)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # 设置页边距
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.17)
        section.right_margin = Cm(3.17)

    # ========== 封面 ==========
    for _ in range(4):
        doc.add_paragraph()

    # 软件名称
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('计算机软件操作手册')
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.name = '黑体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('企业资产与物资管理系统 V1.0')
    run.font.size = Pt(18)
    run.font.name = '黑体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

    for _ in range(6):
        doc.add_paragraph()

    # 封面信息
    cover_info = [
        ('著作权人：', '________________'),
        ('版 本 号：', 'V1.0'),
        ('编写日期：', '2026年05月'),
    ]

    for label, value in cover_info:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(label)
        run.font.size = Pt(14)
        run.font.name = '宋体'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        run = p.add_run(value)
        run.font.size = Pt(14)
        run.font.name = '宋体'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    doc.add_page_break()

    # ========== 目录页 ==========
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('目  录')
    run.font.size = Pt(18)
    run.font.bold = True
    run.font.name = '黑体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

    doc.add_paragraph()

    # 目录内容
    toc_items = [
        '1 引言',
        '　1.1 编写目的',
        '　1.2 项目背景',
        '　1.3 术语与缩略语',
        '　1.4 参考资料',
        '',
        '2 软件概述',
        '　2.1 软件用途',
        '　2.2 软件功能',
        '　2.3 软件特点',
        '',
        '3 运行环境',
        '　3.1 硬件环境',
        '　3.2 软件环境',
        '　3.3 网络环境',
        '',
        '4 软件安装与部署',
        '　4.1 服务端部署',
        '　4.2 前端部署',
        '　4.3 微信小程序部署',
        '　4.4 数据库初始化',
        '',
        '5 软件操作说明',
        '　5.1 系统登录',
        '　5.2 系统首页',
        '　5.3 系统管理',
        '　5.4 仓库管理',
        '　5.5 物资入库管理',
        '　5.6 物资管理',
        '　5.7 资产管理',
        '　5.8 巡检管理',
        '　5.9 报修管理',
        '　5.10 使用说明书管理',
        '　5.11 系统监控',
        '　5.12 系统工具',
        '　5.13 微信小程序端',
        '',
        '6 异常处理与常见问题',
        '',
        '7 附录',
    ]

    for item in toc_items:
        if item:
            p = doc.add_paragraph(item)
            p.paragraph_format.line_spacing = 1.5

    doc.add_page_break()

    # ========== 正文内容 ==========
    def add_heading(text, level=1):
        p = doc.add_paragraph()
        run = p.add_run(text)
        if level == 1:
            run.font.size = Pt(18)
            run.font.bold = True
            run.font.name = '黑体'
            run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(6)
        elif level == 2:
            run.font.size = Pt(15)
            run.font.bold = True
            run.font.name = '黑体'
            run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(6)
        elif level == 3:
            run.font.size = Pt(14)
            run.font.bold = True
            run.font.name = '黑体'
            run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(4)

    def add_paragraph(text, bold=False):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.size = Pt(12)
        run.font.name = '宋体'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        if bold:
            run.font.bold = True
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.first_line_indent = Cm(0.74)
        return p

    def add_table(headers, rows):
        table = doc.add_table(rows=1 + len(rows), cols=len(headers))
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # 表头
        for i, header in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = header
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_cell_shading(cell, 'D9E2F3')
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.size = Pt(11)
                    run.font.name = '宋体'
                    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

        # 数据行
        for i, row in enumerate(rows):
            for j, cell_text in enumerate(row):
                cell = table.rows[i + 1].cells[j]
                cell.text = cell_text
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(11)
                        run.font.name = '宋体'
                        run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

        doc.add_paragraph()

    # ========== 第1章 引言 ==========
    add_heading('1 引言', 1)

    add_heading('1.1 编写目的', 2)
    add_paragraph('本手册旨在为"企业资产与物资管理系统"的使用者提供详细的操作指导，帮助用户快速掌握系统的各项功能，确保系统能够被正确、高效地使用。本手册适用于系统管理员、仓库管理员、巡检人员、维修人员及其他业务操作人员。')

    add_heading('1.2 项目背景', 2)
    add_paragraph('随着企业规模的不断扩大，资产与物资的管理日趋复杂，传统的人工管理方式已难以满足现代化管理需求。为提高企业资产与物资管理的效率和准确性，降低管理成本，特开发本系统。本系统基于成熟的若依（RuoYi）快速开发框架构建，采用前后端分离架构，实现了资产全生命周期管理，涵盖物资入库、出库、巡检、报修、报废等全流程管理。')

    add_heading('1.3 术语与缩略语', 2)
    add_table(
        ['术语/缩略语', '说明'],
        [
            ['ERP', '企业资源计划（Enterprise Resource Planning）'],
            ['JWT', 'JSON Web Token，用于身份认证的令牌机制'],
            ['API', '应用程序编程接口（Application Programming Interface）'],
            ['CRUD', '增删改查（Create、Read、Update、Delete）'],
            ['ORM', '对象关系映射（Object Relational Mapping）'],
            ['RESTful', '一种API设计风格'],
        ]
    )

    add_heading('1.4 参考资料', 2)
    refs = [
        '若依官方文档：http://doc.ruoyi.vip',
        'Spring Boot官方文档',
        'Vue.js官方文档',
        'Element UI组件库文档',
    ]
    for ref in refs:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(ref)
        run.font.size = Pt(12)
        run.font.name = '宋体'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    doc.add_page_break()

    # ========== 第2章 软件概述 ==========
    add_heading('2 软件概述', 1)

    add_heading('2.1 软件用途', 2)
    add_paragraph('本系统是一套面向企业的资产与物资综合管理平台，主要用于：')

    purposes = [
        '对企业仓库及物资进行信息化管理，实现物资从采购入库到报废出库的全生命周期跟踪。',
        '管理资产流转记录，确保资产去向清晰、责任明确。',
        '制定并执行设备巡检计划，及时发现设备异常，保障设备正常运行。',
        '提供便捷的报修申请与处理通道，缩短维修响应时间。',
        '管理设备使用说明书等技术文档，方便相关人员查阅。',
        '通过微信小程序实现移动端的巡检、报修等操作，提高现场工作效率。',
    ]
    for i, purpose in enumerate(purposes, 1):
        add_paragraph(f'{i}. {purpose}')

    add_heading('2.2 软件功能', 2)
    add_paragraph('本系统主要包含以下功能模块：')
    add_table(
        ['序号', '功能模块', '功能说明'],
        [
            ['1', '系统管理', '用户管理、部门管理、岗位管理、角色管理、菜单管理、字典管理、参数设置、通知公告'],
            ['2', '仓库管理', '仓库信息的增删改查，支持仓库编码、地址、联系人等信息维护'],
            ['3', '物资入库', '入库单管理，支持一单多物资入库，含供应商、发票信息、明细管理及审核流程'],
            ['4', '物资管理', '物资分类管理、物资信息管理，支持物资编码、规格型号、库存数量、单价、保修期限等信息维护'],
            ['5', '资产管理', '资产流转记录（入库、出库、报损、报废）、资产变更记录'],
            ['6', '巡检管理', '巡检项目管理、巡检记录管理，支持按周期制定巡检计划'],
            ['7', '报修管理', '报修记录的提交、处理、完成、驳回、取消等全流程管理'],
            ['8', '使用说明书管理', '设备使用说明书的上传、关联、版本管理'],
            ['9', '系统监控', '操作日志、登录日志、在线用户、定时任务、服务监控、缓存监控、数据监控'],
            ['10', '系统工具', '代码生成、系统接口文档'],
            ['11', '微信小程序', '移动端登录、资产查看、巡检报告、报修申请与处理、个人中心'],
        ]
    )

    add_heading('2.3 软件特点', 2)
    features = [
        ('前后端分离架构：', '后端采用Spring Boot框架，前端采用Vue.js + Element UI，松耦合设计便于独立开发与部署。'),
        ('多终端支持：', '除Web端外，提供微信小程序端，支持现场扫码巡检、报修等移动办公场景。'),
        ('权限控制精细：', '基于RBAC权限模型，支持菜单权限、按钮权限、数据权限的多级控制。'),
        ('数据安全可靠：', '采用JWT令牌认证、XSS防护、SQL注入防护等安全机制，保障系统数据安全。'),
        ('操作日志完整：', '系统自动记录用户操作日志和登录日志，便于追溯与审计。'),
        ('代码生成高效：', '内置代码生成器，可一键生成前后端CRUD代码，提高开发效率。'),
    ]
    for i, (title, desc) in enumerate(features, 1):
        p = doc.add_paragraph()
        run = p.add_run(f'{i}. {title}')
        run.font.bold = True
        run.font.size = Pt(12)
        run.font.name = '宋体'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        run = p.add_run(desc)
        run.font.size = Pt(12)
        run.font.name = '宋体'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.first_line_indent = Cm(0.74)

    doc.add_page_break()

    # ========== 第3章 运行环境 ==========
    add_heading('3 运行环境', 1)

    add_heading('3.1 硬件环境', 2)

    p = add_paragraph('服务器端：', bold=True)
    p.paragraph_format.first_line_indent = Cm(0)

    add_table(
        ['项目', '最低配置', '推荐配置'],
        [
            ['CPU', '2核', '4核及以上'],
            ['内存', '4GB', '8GB及以上'],
            ['硬盘', '50GB', '100GB及以上'],
            ['网络', '100Mbps', '1000Mbps'],
        ]
    )

    p = add_paragraph('客户端：', bold=True)
    p.paragraph_format.first_line_indent = Cm(0)

    add_table(
        ['项目', '最低配置'],
        [
            ['CPU', '1GHz'],
            ['内存', '2GB'],
            ['显示器', '分辨率1024×768及以上'],
            ['网络', '10Mbps'],
        ]
    )

    add_heading('3.2 软件环境', 2)

    p = add_paragraph('服务端：', bold=True)
    p.paragraph_format.first_line_indent = Cm(0)

    add_table(
        ['类型', '名称', '版本'],
        [
            ['操作系统', 'Windows Server / Linux（CentOS/Ubuntu）', '不限'],
            ['JDK', 'Java Development Kit', '1.8及以上'],
            ['数据库', 'MySQL', '5.7及以上'],
            ['缓存', 'Redis', '3.0及以上'],
            ['Web服务器', 'Apache Tomcat（内嵌）', '9.0'],
            ['应用框架', 'Spring Boot', '2.5.15'],
            ['安全框架', 'Spring Security', '5.7.14'],
            ['ORM框架', 'MyBatis', '不限'],
            ['连接池', 'Druid', '1.2.28'],
            ['接口文档', 'Swagger', '3.0.0'],
        ]
    )

    p = add_paragraph('客户端：', bold=True)
    p.paragraph_format.first_line_indent = Cm(0)

    add_table(
        ['类型', '名称', '版本'],
        [
            ['浏览器', 'Chrome / Firefox / Edge / Safari', '最新稳定版'],
            ['前端框架', 'Vue.js', '2.x'],
            ['UI组件库', 'Element UI', '不限'],
            ['微信小程序', '微信开发者工具', '最新稳定版'],
        ]
    )

    add_heading('3.3 网络环境', 2)
    network_items = [
        '服务器需部署在可访问的网络环境中，确保客户端能够通过HTTP/HTTPS协议访问。',
        '服务器需开放8080端口（Web服务端口）。',
        '数据库服务器（MySQL）需开放3306端口。',
        '缓存服务器（Redis）需开放6379端口。',
        '微信小程序端需确保服务器域名已配置HTTPS证书并在微信公众平台完成域名白名单配置。',
    ]
    for i, item in enumerate(network_items, 1):
        add_paragraph(f'{i}. {item}')

    doc.add_page_break()

    # ========== 第4章 软件安装与部署 ==========
    add_heading('4 软件安装与部署', 1)

    add_heading('4.1 服务端部署', 2)

    steps = [
        ('步骤一：安装JDK环境', [
            '下载并安装JDK 1.8或以上版本。',
            '配置JAVA_HOME环境变量，将其指向JDK安装目录。',
            '将%JAVA_HOME%/bin添加到系统PATH环境变量中。',
            '执行java -version命令验证安装是否成功。',
        ]),
        ('步骤二：安装MySQL数据库', [
            '下载并安装MySQL 5.7或以上版本。',
            '创建数据库，执行sql/目录下的初始化脚本完成数据库表的创建和初始数据的导入。',
            '修改ruoyi-admin/src/main/resources/application-druid.yml中的数据库连接信息。',
        ]),
        ('步骤三：安装Redis缓存', [
            '下载并安装Redis 3.0或以上版本。',
            '启动Redis服务。',
            '如需配置密码，修改ruoyi-admin/src/main/resources/application.yml中的Redis连接配置。',
        ]),
        ('步骤四：启动后端服务', [
            '使用Maven构建项目：在项目根目录执行mvn clean package命令。',
            '运行打包后的jar文件：java -jar ruoyi-admin.jar。',
            '或使用开发工具（如IntelliJ IDEA）直接运行RuoYiApplication主类启动服务。',
            '服务启动成功后，默认监听8080端口。',
        ]),
    ]

    for step_title, step_items in steps:
        p = add_paragraph(step_title, bold=True)
        p.paragraph_format.first_line_indent = Cm(0)
        for i, item in enumerate(step_items, 1):
            add_paragraph(f'{i}. {item}')

    add_heading('4.2 前端部署', 2)
    steps = [
        ('步骤一：安装Node.js环境', [
            '下载并安装Node.js（建议版本12.x或以上）。',
            '执行node -v和npm -v命令验证安装是否成功。',
        ]),
        ('步骤二：安装前端依赖', [
            '进入ruoyi-ui目录。',
            '执行npm install命令安装项目依赖。',
        ]),
        ('步骤三：配置后端接口地址', [
            '修改ruoyi-ui目录下的接口配置文件，将后端API地址指向实际的服务端地址。',
        ]),
        ('步骤四：启动前端服务', [
            '执行npm run dev命令启动前端开发服务器。',
            '启动成功后，通过浏览器访问http://localhost即可打开系统登录页面。',
        ]),
        ('步骤五：生产环境部署', [
            '执行npm run build命令构建生产环境前端资源。',
            '将生成的dist目录下的文件部署到Nginx或其他Web服务器中。',
        ]),
    ]
    for step_title, step_items in steps:
        p = add_paragraph(step_title, bold=True)
        p.paragraph_format.first_line_indent = Cm(0)
        for i, item in enumerate(step_items, 1):
            add_paragraph(f'{i}. {item}')

    add_heading('4.3 微信小程序部署', 2)
    mp_steps = [
        '下载并安装微信开发者工具。',
        '使用微信开发者工具打开mini-program目录。',
        '在mini-program/utils/request.js中配置后端API接口地址。',
        '在微信公众平台完成小程序的注册、审核和发布。',
    ]
    for i, item in enumerate(mp_steps, 1):
        add_paragraph(f'{i}. {item}')

    add_heading('4.4 数据库初始化', 2)
    db_steps = [
        '创建MySQL数据库，字符集设置为utf8mb4。',
        '按顺序执行sql/目录下的SQL脚本：ry_20260417.sql（基础系统表结构和数据）、biz_tables.sql（业务模块表结构）、stock_in.sql（物资入库功能表结构）。',
        '验证数据库表是否创建成功。',
    ]
    for i, item in enumerate(db_steps, 1):
        add_paragraph(f'{i}. {item}')

    doc.add_page_break()

    # ========== 第5章 软件操作说明 ==========
    add_heading('5 软件操作说明', 1)

    add_heading('5.1 系统登录', 2)
    add_paragraph('用户通过浏览器访问系统地址，进入系统登录页面。', bold=True)

    p = add_paragraph('操作步骤：', bold=True)
    p.paragraph_format.first_line_indent = Cm(0)

    login_steps = [
        '打开浏览器，在地址栏输入系统访问地址（如：http://服务器IP:8080）。',
        '系统显示登录页面，包含用户名输入框、密码输入框和验证码输入框。',
        '在"用户名"输入框中输入分配的用户账号。',
        '在"密码"输入框中输入对应的密码。',
        '在"验证码"输入框中输入页面显示的数学计算验证码结果。',
        '点击"登录"按钮。',
        '系统验证通过后自动跳转至系统首页。',
    ]
    for i, step in enumerate(login_steps, 1):
        add_paragraph(f'{i}. {step}')

    p = add_paragraph('注意事项：', bold=True)
    p.paragraph_format.first_line_indent = Cm(0)

    notes = [
        '连续输入错误密码5次，账号将被锁定10分钟。',
        '验证码不区分大小写。',
        '首次登录建议修改初始密码。',
    ]
    for note in notes:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(note)
        run.font.size = Pt(12)
        run.font.name = '宋体'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    add_heading('5.2 系统首页', 2)
    add_paragraph('登录成功后进入系统首页，首页展示系统导航菜单和常用功能入口。', bold=True)

    features = [
        '顶部导航栏：显示系统名称、全屏切换、布局大小设置、用户头像及下拉菜单。',
        '左侧菜单栏：以树形结构展示系统所有功能模块菜单，点击菜单项可进入对应功能页面。',
        '主内容区：显示当前操作页面的内容。',
        '标签页导航：已打开的页面以标签页形式展示，方便在多个页面间快速切换。',
    ]
    for feature in features:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(feature)
        run.font.size = Pt(12)
        run.font.name = '宋体'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    add_heading('5.3 系统管理', 2)

    # 5.3.1 用户管理
    add_heading('5.3.1 用户管理', 3)
    add_paragraph('管理系统用户信息，包括新增、修改、删除、重置密码、分配角色等操作。', bold=True)

    user_steps = [
        '在左侧菜单中点击"系统管理" -> "用户管理"，进入用户管理页面。',
        '页面以表格形式展示所有用户信息，支持按用户名、手机号、状态等条件进行筛选查询。',
        '新增用户：点击"新增"按钮，弹出新增用户对话框，填写用户信息，点击"确定"保存。',
        '修改用户：点击目标用户所在行的"修改"按钮，修改相应信息后点击"确定"保存。',
        '删除用户：勾选目标用户，点击"删除"按钮，确认后删除。',
        '重置密码：点击目标用户所在行的"重置密码"按钮，可将该用户密码重置为默认密码。',
        '导出数据：点击"导出"按钮，可将当前查询结果导出为Excel文件。',
    ]
    for i, step in enumerate(user_steps, 1):
        add_paragraph(f'{i}. {step}')

    # 5.3.2-5.3.8 其他系统管理功能
    other_sys_funcs = [
        ('5.3.2 部门管理', '配置系统组织机构（公司、部门、小组），以树形结构展现，支持数据权限控制。'),
        ('5.3.3 岗位管理', '配置系统用户所属担任职务。'),
        ('5.3.4 角色管理', '管理系统角色，配置角色的菜单权限和数据权限。'),
        ('5.3.5 菜单管理', '配置系统菜单、操作权限和按钮权限标识。'),
        ('5.3.6 字典管理', '对系统中经常使用的一些较为固定的数据进行维护，如状态、类型等枚举值。'),
        ('5.3.7 参数设置', '对系统动态配置常用参数。'),
        ('5.3.8 通知公告', '系统通知公告信息的发布与维护。'),
    ]
    for title, desc in other_sys_funcs:
        add_heading(title, 3)
        add_paragraph(desc, bold=True)

    add_heading('5.4 仓库管理', 2)
    add_heading('5.4.1 仓库信息管理', 3)
    add_paragraph('管理企业仓库的基本信息，包括仓库编码、名称、地址、联系人等。', bold=True)

    warehouse_steps = [
        '在左侧菜单中点击"仓库管理" -> "仓库信息管理"，进入仓库信息管理页面。',
        '页面以表格形式展示所有仓库信息，支持按仓库编码、仓库名称、状态等条件进行筛选查询。',
        '新增仓库：点击"新增"按钮，填写仓库编码、名称、地址、联系人等信息，点击"确定"保存。',
        '修改仓库：点击目标仓库所在行的"修改"按钮，修改信息后保存。',
        '删除仓库：勾选目标仓库，点击"删除"按钮，确认后删除。',
        '导出数据：点击"导出"按钮，可将仓库信息导出为Excel文件。',
    ]
    for i, step in enumerate(warehouse_steps, 1):
        add_paragraph(f'{i}. {step}')

    add_heading('5.5 物资入库管理', 2)
    add_heading('5.5.1 入库单列表', 3)
    add_paragraph('管理物资入库单，支持一单多物资入库，共享供应商和发票信息，实现入库全流程管理。', bold=True)
    stock_in_steps = [
        '在左侧菜单中点击"仓库管理" -> "物资入库"，进入入库单列表页面。',
        '页面以表格形式展示所有入库单，支持按入库单号、供应商、入库仓库、状态、入库日期等条件筛选查询。',
        '状态说明：待审核（蓝色标签）、已审核（绿色标签）、已驳回（红色标签）。',
        '操作按钮：详情、修改（待审核状态）、审核（待审核状态）、删除（待审核状态）。',
        '导出数据：点击"导出"按钮，可将入库单信息导出为Excel文件。',
    ]
    for i, step in enumerate(stock_in_steps, 1):
        add_paragraph(f'{i}. {step}')

    add_heading('5.5.2 新增入库单', 3)
    add_paragraph('点击"新增入库"按钮，填写主表信息（供应商、入库仓库、入库日期、发票信息等），然后添加入库明细。', bold=True)
    stock_in_form_steps = [
        '入库单号：系统自动生成，无需手动输入。',
        '供应商：输入供应商名称（必填）。',
        '入库仓库：从下拉列表中选择入库目标仓库（必填）。',
        '入库日期：选择入库日期（必填）。',
        '点击"添加物资"按钮，从下拉列表中选择已有物资，自动填充规格型号、单位、单价、保修期限等信息。',
        '支持添加多条明细记录，系统自动计算入库总数量和总金额。',
        '点击"确定"保存入库单。',
    ]
    for i, step in enumerate(stock_in_form_steps, 1):
        add_paragraph(f'{i}. {step}')

    add_heading('5.5.3 入库单详情', 3)
    add_paragraph('点击"详情"按钮查看入库单完整信息，包括主表信息和入库明细。', bold=True)
    add_paragraph('入库单审核流程：新增默认"待审核"，审核人可选择"通过"或"驳回"，已审核的入库单不允许修改和删除。')

    add_heading('5.6 物资管理', 2)
    add_heading('5.6.1 物资分类管理', 3)
    add_paragraph('管理物资的分类信息，支持多级分类。', bold=True)

    add_heading('5.6.2 物资信息管理', 3)
    add_paragraph('管理企业物资的详细信息，包括物资编码、名称、分类、规格型号、库存数量、单价、保修期限、供应商等。', bold=True)

    add_heading('5.7 资产管理', 2)
    add_heading('5.7.1 资产流转记录', 3)
    add_paragraph('记录物资的入库、出库、报损、报废等流转操作，实现资产去向的全程跟踪。', bold=True)

    add_heading('5.7.2 资产变更记录', 3)
    add_paragraph('记录资产的位置变更、科室变更、状态变更等信息。', bold=True)

    add_heading('5.8 巡检管理', 2)
    add_heading('5.8.1 巡检项目管理', 3)
    add_paragraph('管理巡检检查项目，定义巡检需要检查的具体内容。', bold=True)

    add_heading('5.8.2 巡检记录管理', 3)
    add_paragraph('记录设备巡检的执行情况，包括巡检人、巡检时间、巡检结果、巡检照片等。', bold=True)

    add_heading('5.9 报修管理', 2)
    add_heading('5.9.1 报修记录管理', 3)
    add_paragraph('管理设备报修的全流程，包括报修提交、处理指派、维修记录、完成确认等环节。', bold=True)

    add_paragraph('报修状态流转：待处理 → 处理中 → 已完成；可驳回或取消。')

    add_heading('5.10 使用说明书管理', 2)
    add_paragraph('管理设备使用说明书等技术文档，支持文档的上传、下载、版本管理，可关联到仓库或物资。', bold=True)

    add_heading('5.11 系统监控', 2)
    monitor_funcs = [
        ('5.11.1 操作日志', '记录系统用户的所有操作行为，支持按操作模块、操作人员、操作时间等条件查询。'),
        ('5.11.2 登录日志', '记录系统用户的登录行为，包括登录成功和登录失败的记录。'),
        ('5.11.3 在线用户', '监控当前系统中的活跃用户状态。'),
        ('5.11.4 定时任务', '管理系统的定时调度任务，支持在线添加、修改、删除任务。'),
        ('5.11.5 服务监控', '监视当前系统的CPU、内存、磁盘、JVM堆栈等服务器资源使用情况。'),
        ('5.11.6 缓存监控', '查看系统的Redis缓存信息，包括缓存命令统计、内存使用、Key数量等。'),
        ('5.11.7 数据监控', '监视当前系统数据库连接池状态，可进行SQL分析找出系统性能瓶颈。'),
    ]
    for title, desc in monitor_funcs:
        add_heading(title, 3)
        add_paragraph(desc, bold=True)

    add_heading('5.12 系统工具', 2)
    add_heading('5.12.1 代码生成', 3)
    add_paragraph('根据数据库表结构一键生成前后端CRUD代码（Java、HTML、XML、SQL），支持自定义模板配置。', bold=True)

    add_heading('5.12.2 系统接口', 3)
    add_paragraph('根据业务代码自动生成API接口文档（基于Swagger），方便前后端联调。', bold=True)

    add_heading('5.13 微信小程序端', 2)

    add_heading('5.13.1 小程序登录', 3)
    add_paragraph('用户通过微信扫码或搜索小程序名称进入系统。', bold=True)
    mp_login_steps = [
        '打开微信，扫描小程序二维码或在微信中搜索小程序名称。',
        '进入小程序登录页面。',
        '输入用户名和密码进行登录。',
        '登录成功后进入小程序首页。',
    ]
    for i, step in enumerate(mp_login_steps, 1):
        add_paragraph(f'{i}. {step}')

    add_heading('5.13.2 首页功能', 3)
    add_paragraph('小程序首页提供常用功能的快捷入口。', bold=True)

    add_heading('5.13.3 资产详情查看', 3)
    add_paragraph('用户可通过扫码或搜索方式查看资产详细信息。', bold=True)

    add_heading('5.13.4 巡检报告', 3)
    add_paragraph('巡检人员通过小程序提交巡检报告。', bold=True)

    add_heading('5.13.5 报修申请', 3)
    add_paragraph('用户通过小程序提交设备报修申请。', bold=True)

    add_heading('5.13.6 报修处理', 3)
    add_paragraph('维修人员通过小程序处理报修工单。', bold=True)

    add_heading('5.13.7 个人中心', 3)
    add_paragraph('用户在个人中心查看和管理个人信息。', bold=True)

    doc.add_page_break()

    # ========== 第6章 异常处理与常见问题 ==========
    add_heading('6 异常处理与常见问题', 1)

    add_heading('6.1 登录异常', 2)
    add_table(
        ['异常现象', '可能原因', '解决方法'],
        [
            ['提示"用户名或密码错误"', '输入的用户名或密码不正确', '检查用户名和密码是否正确，注意大小写'],
            ['提示"验证码错误"', '验证码输入错误', '重新输入验证码，注意数学计算结果'],
            ['提示"用户已被锁定"', '连续5次输入错误密码', '等待10分钟后重试，或联系管理员解锁'],
            ['登录页面无法访问', '服务未启动或网络异常', '检查服务是否正常运行，检查网络连接'],
        ]
    )

    add_heading('6.2 数据操作异常', 2)
    add_table(
        ['异常现象', '可能原因', '解决方法'],
        [
            ['新增/修改数据失败', '必填项未填写或格式不正确', '检查表单中带*号的必填项是否已填写'],
            ['删除数据失败', '该数据被其他模块引用', '先删除关联数据后再删除目标数据'],
            ['查询结果为空', '查询条件过于严格或数据不存在', '放宽查询条件或确认数据是否存在'],
            ['导出Excel失败', '数据量过大或服务器异常', '减少导出数据量，或稍后重试'],
        ]
    )

    add_heading('6.3 网络异常', 2)
    add_table(
        ['异常现象', '可能原因', '解决方法'],
        [
            ['页面加载缓慢', '网络延迟或服务器负载高', '检查网络状况，或联系管理员检查服务器状态'],
            ['接口请求超时', '服务器响应超时', '检查后端服务是否正常，检查数据库连接'],
            ['页面显示"网络错误"', '网络连接中断', '检查网络连接，刷新页面重试'],
        ]
    )

    add_heading('6.4 文件上传异常', 2)
    add_table(
        ['异常现象', '可能原因', '解决方法'],
        [
            ['文件上传失败', '文件大小超过限制', '单个文件最大支持10MB，请压缩后重新上传'],
            ['图片上传失败', '图片格式不支持', '上传JPG、PNG、GIF等常见图片格式'],
            ['说明书上传失败', '文件格式不支持', '上传PDF、Word等支持的文档格式'],
        ]
    )

    doc.add_page_break()

    # ========== 第7章 附录 ==========
    add_heading('7 附录', 1)

    add_heading('7.1 数据字典说明', 2)

    # 报修优先级
    p = add_paragraph('报修优先级：', bold=True)
    p.paragraph_format.first_line_indent = Cm(0)
    add_table(
        ['字典值', '显示标签', '说明'],
        [
            ['low', '低', '非紧急情况，可安排常规维修'],
            ['normal', '普通', '一般故障，按正常流程处理'],
            ['high', '高', '较为紧急，需优先处理'],
            ['urgent', '紧急', '紧急故障，需立即处理'],
        ]
    )

    # 报修状态
    p = add_paragraph('报修状态：', bold=True)
    p.paragraph_format.first_line_indent = Cm(0)
    add_table(
        ['字典值', '显示标签', '说明'],
        [
            ['pending', '待处理', '报修已提交，等待处理'],
            ['processing', '处理中', '已指派维修人员，正在维修'],
            ['completed', '已完成', '维修已完成'],
            ['rejected', '已驳回', '报修被驳回'],
            ['cancelled', '已取消', '报修已取消'],
        ]
    )

    # 流转类型
    p = add_paragraph('流转类型：', bold=True)
    p.paragraph_format.first_line_indent = Cm(0)
    add_table(
        ['字典值', '显示标签', '说明'],
        [
            ['IN', '入库', '物资入库存放'],
            ['OUT', '出库', '物资出库使用'],
            ['DAMAGE', '报损', '物资损坏报告'],
            ['SCRAP', '报废', '物资报废处理'],
        ]
    )

    # 巡检周期
    p = add_paragraph('巡检周期：', bold=True)
    p.paragraph_format.first_line_indent = Cm(0)
    add_table(
        ['字典值', '显示标签', '说明'],
        [
            ['daily', '每日', '每日巡检'],
            ['weekly', '每周', '每周巡检'],
            ['monthly', '每月', '每月巡检'],
            ['quarterly', '每季度', '每季度巡检'],
            ['yearly', '每年', '每年巡检'],
        ]
    )

    add_heading('7.2 系统权限说明', 2)
    add_paragraph('系统采用基于RBAC（Role-Based Access Control）的权限控制模型，权限粒度分为以下层级：')

    p = add_paragraph('菜单权限：', bold=True)
    p.paragraph_format.first_line_indent = Cm(0)
    add_paragraph('控制用户可访问的菜单和页面。')

    p = add_paragraph('按钮权限：', bold=True)
    p.paragraph_format.first_line_indent = Cm(0)
    add_paragraph('控制页面中各操作按钮（新增、修改、删除、导出等）的可见性。')

    p = add_paragraph('数据权限：', bold=True)
    p.paragraph_format.first_line_indent = Cm(0)
    add_paragraph('控制用户可查看的数据范围。')

    add_table(
        ['数据权限范围', '说明'],
        [
            ['全部数据权限', '可查看系统所有数据'],
            ['自定义数据权限', '可查看指定部门的数据'],
            ['本部门数据权限', '仅可查看本部门数据'],
            ['本部门及以下数据权限', '可查看本部门及下级部门数据'],
            ['仅本人数据权限', '仅可查看本人创建的数据'],
        ]
    )

    add_heading('7.3 快捷操作说明', 2)
    add_table(
        ['操作', '说明'],
        [
            ['刷新页面', '点击标签页上的刷新按钮'],
            ['全屏显示', '点击顶部导航栏的全屏按钮'],
            ['关闭标签页', '点击标签页上的关闭按钮'],
            ['关闭其他标签页', '右键点击标签页，选择"关闭其他"'],
            ['关闭所有标签页', '右键点击标签页，选择"关闭所有"'],
            ['修改密码', '点击右上角头像 -> 个人中心 -> 修改密码'],
            ['清空查询条件', '点击查询区域的"重置"按钮'],
        ]
    )

    # 结束语
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('文档编写完成')
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.name = '黑体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('本手册详细介绍了"企业资产与物资管理系统V1.0"的安装部署、功能操作、异常处理等内容。如有疑问，请联系系统管理员。')
    run.font.size = Pt(12)
    run.font.name = '宋体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    return doc

if __name__ == '__main__':
    doc = create_document()
    output_path = os.path.join(os.path.dirname(__file__), '计算机软件操作手册.docx')
    doc.save(output_path)
    print(f'Word文档已生成: {output_path}')
